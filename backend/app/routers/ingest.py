"""
Manual Ingestion Router — /ingest/*

Two-step workflow:
  POST /ingest/raw     → Claude parses pasted text → returns preview
  POST /ingest/approve → validate → insert to Supabase → return id

Completely separate from the scraping pipeline.
"""
import io
import json
from typing import Any, Optional
from fastapi import APIRouter, HTTPException, UploadFile, File
from pydantic import BaseModel
from app.database.supabase_client import get_supabase
from app.utils.claude_utils import get_claude, ClaudeEmptyResponseError, ClaudeOverloadedError
from app.utils.name_utils import normalize_name, find_fighter_by_name
from app.utils.ingest_validator import validate_fighter_profile, validate_fight_history
from app.utils.openai_audit import audit_parse
from app.prompts import RAW_FIGHTER_PROFILE_PROMPT, RAW_FIGHT_HISTORY_PROMPT

router = APIRouter(prefix="/ingest", tags=["Manual Ingestion"])


# ── Request / Response models ──────────────────────────────────────────────────

class RawIngestRequest(BaseModel):
    type: str                      # "fighter_profile" | "fight_history"
    raw_text: str
    fighter_name: Optional[str] = None   # required for fight_history


class ApproveRequest(BaseModel):
    type: str
    data: Any                      # dict (fighter_profile) or list (fight_history)
    fighter_id: Optional[str] = None    # required for fight_history


# ── Helpers ────────────────────────────────────────────────────────────────────

def _performance_from(data: dict) -> dict:
    def _i(v, d=0): return int(v) if isinstance(v, (int, float)) else d
    def _f(v, d=None): return float(v) if isinstance(v, (int, float)) else d
    wins = _i(data.get("wins"))
    total = wins + _i(data.get("losses"))
    ko  = _i(data.get("ko_wins"))
    tko = _i(data.get("tko_wins"))
    sub = _i(data.get("sub_wins"))
    finish_rate = round((ko + tko + sub) / total, 3) if total > 0 else None
    return {
        "ko_wins":                  ko,
        "tko_wins":                 tko,
        "submission_wins":          sub,
        "decision_wins":            _i(data.get("dec_wins")),
        "finish_rate":              finish_rate,
        "losses_by_ko":             _i(data.get("losses_by_ko")),
        "losses_by_submission":     _i(data.get("losses_by_submission")),
        "losses_by_decision":       _i(data.get("losses_by_decision")),
        "strikes_landed_per_min":   _f(data.get("slpm")),
        "strikes_absorbed_per_min": _f(data.get("sapm")),
        "strike_accuracy":          _f(data.get("strike_accuracy")),
        "strike_defense":           _f(data.get("strike_defense")),
        "takedown_avg":             _f(data.get("takedown_avg")),
        "takedown_accuracy":        _f(data.get("takedown_accuracy")),
        "takedown_defense":         _f(data.get("takedown_defense")),
        "submission_avg":           _f(data.get("submission_avg")),
        "longest_win_streak":       _i(data.get("longest_win_streak")),
    }


def _normalize_event_name(name: str) -> str:
    """
    Standardize event names for deduplication.
    e.g. "UFC 300: Pereira vs Hill" -> "UFC 300"
    """
    if not name: return ""
    name = name.strip()
    # Strip "UFC XXX: ..." or "UFC Fight Night: ..." prefixes
    import re
    # Match "UFC <digits>:" or "UFC Fight Night <digits / name>:"
    # We want to keep the "UFC <digits>" or "UFC Fight Night <something>"
    match = re.match(r"^(UFC\s+(?:\d+|Fight\s+Night[^\:]+))\:", name, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return name

def _resolve_or_create_opponent(supabase, opponent_name: str) -> tuple:
    """
    Given an opponent name from a bout, return (fighter_id, was_created).

    Strategy:
      1. Normalize the name and look it up with find_fighter_by_name (3-stage:
         exact → alias → fuzzy).  If found → return existing ID, was_created=False.
      2. If not found → insert a minimal shell fighter record so the fight_history
         row has a proper FK.  Returns (new_id, True).
      3. On any error → returns (None, False).  The row is still inserted without
         opponent_id rather than blocking the whole approve.

    NEVER re-parses or re-ingests a fighter that already exists.
    """
    if not opponent_name.strip():
        return None, False

    opp_first, opp_last = normalize_name(opponent_name)
    if not opp_first:
        return None, False

    # Stage 1 — look up existing fighter
    try:
        existing = find_fighter_by_name(supabase, opp_first, opp_last)
        if existing:
            print(f"[Ingest] Opponent '{opponent_name}' found in DB ({existing['id']})")
            return existing["id"], False
    except Exception as e:
        print(f"[Ingest] Opponent lookup error for '{opponent_name}': {e}")
        return None, False

    # Stage 2 — create minimal shell so the FK resolves
    try:
        first_cap = opp_first.capitalize()
        last_cap  = " ".join(w.capitalize() for w in opp_last.split()) if opp_last else ""
        shell = {
            "first_name":    first_cap,
            "last_name":     last_cap,
            "gender":        "Male",
            "nationality":   "",
            "weight_class":  "Unknown",
            "organization":  "Unknown",
            "stance":        "",
            "gym":           "",
            "head_coach":    "",
            "image_url":     "",
            "notes":         f"Shell — auto-created during fight history ingestion (opponent: {opponent_name})",
            "is_active":     True,
            "is_champion":   False,
            "is_verified":   False,
            "admin_status":  "pending",
            "wins":          0,
            "losses":        0,
            "draws":         0,
            "nc":            0,
        }
        res = supabase.table("fighters").insert(shell).execute()
        new_id = res.data[0]["id"]
        print(f"[Ingest] Created shell fighter '{opponent_name}' → {new_id}")
        return new_id, True
    except Exception as e:
        print(f"[Ingest] Failed to create shell for '{opponent_name}': {e}")
        return None, False


def _build_fighter_upsert(data: dict) -> dict:
    """Convert preview data dict → DB upsert payload."""
    first_name = str(data.get("first_name") or "").strip()
    last_name  = str(data.get("last_name") or "").strip()
    # Normalise capitalisation
    first_name = " ".join(w.capitalize() for w in first_name.split())
    last_name  = " ".join(w.capitalize() for w in last_name.split())

    stance = data.get("stance")
    if stance not in ("Orthodox", "Southpaw", "Switch"):
        stance = ""   # empty string avoids NOT NULL violation; filtered at push time

    perf = _performance_from(data)

    upsert: dict = {
        "first_name":      first_name,
        "last_name":       last_name,
        "nickname":        data.get("nickname") or None,
        "gender":          data.get("gender") or "Male",
        "nationality":     data.get("nationality") or "",
        "date_of_birth":   data.get("date_of_birth") or None,
        "organization":    data.get("organization") or "UFC",
        "weight_class":    data.get("weight_class"),
        "height_inch":     data.get("height_inch"),
        "reach_inch":      data.get("reach_inch"),
        "stance":          stance,
        "is_active":       data.get("is_active", True),
        "is_champion":     data.get("is_champion", False),
        "ranking":         data.get("ranking"),
        "gym":             data.get("gym") or data.get("team") or "",
        "head_coach":      data.get("head_coach") or "",
        "fighting_out_of": data.get("fighting_out_of") or None,
        "wins":            int(data.get("wins") or 0),
        "losses":          int(data.get("losses") or 0),
        "draws":           int(data.get("draws") or 0),
        "nc":              int(data.get("nc") or 0),
        "image_url":       data.get("image_url") or "",
        "performance":     perf,
        "notes":           data.get("notes") or "",
        "is_verified":     False,
        "admin_status":    "pending",
        # External identity — populated when source URL is available.
        # Used as a stable dedup key once migration 013 is applied.
        "source_url":      data.get("source_url") or None,
        "source_id":       data.get("source_id") or None,
    }

    # Strip truly optional null fields to avoid NOT NULL violations
    NULLABLE = {
        "nickname", "date_of_birth",
        "reach_inch", "fighting_out_of",
        "ranking",
    }
    return {k: v for k, v in upsert.items() if v is not None or k not in NULLABLE}


# ── Event date enrichment ──────────────────────────────────────────────────────

def _enrich_bout_dates(bouts: list) -> list:
    """
    For any bout where event_date is null but event_name is present,
    attempt to resolve the date from the events table via fuzzy name match.
    """
    # Collect unique event names that need dates
    needs_date = {
        b["event_name"] for b in bouts
        if not b.get("event_date") and b.get("event_name")
    }
    if not needs_date:
        return bouts

    supabase = get_supabase()
    date_cache: dict[str, str] = {}

    for event_name in needs_date:
        norm_name = _normalize_event_name(event_name)
        try:
            # Try exact match first, then normalized ilike
            res = supabase.table("events").select("name,date").ilike(
                "name", f"%{norm_name[:40]}%"
            ).limit(1).execute()
            if res.data:
                raw_date = res.data[0]["date"]
                if raw_date:
                    # Normalise to YYYY-MM-DD
                    date_cache[event_name] = str(raw_date)[:10]
        except Exception:
            pass  # silently skip — date stays null

    if not date_cache:
        return bouts

    enriched = []
    for b in bouts:
        name = b.get("event_name")
        if not b.get("event_date") and name and name in date_cache:
            b = {**b, "event_date": date_cache[name]}
        enriched.append(b)
    return enriched


# ── File upload → text extraction ───────────────────────────────────────────────

def _extract_text_from_docx(raw_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(raw_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    return "\n".join(parts)


@router.post("/extract-file")
async def extract_file(file: UploadFile = File(...)):
    """
    Extract plain text from an uploaded source file so the operator doesn't
    have to copy/paste. Returns the text for the SAME /ingest/raw + /ingest/approve
    review flow — this endpoint never writes to the database.
    """
    filename = file.filename or "upload"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    raw_bytes = await file.read()

    if not raw_bytes:
        raise HTTPException(400, "Uploaded file is empty")

    try:
        if ext == "docx":
            text = _extract_text_from_docx(raw_bytes)
        elif ext in ("txt", "md", "csv"):
            text = raw_bytes.decode("utf-8", errors="replace")
        else:
            raise HTTPException(
                400,
                f"Unsupported file type '.{ext}'. Upload .docx, .txt, .md, or .csv — "
                "for PDFs, copy/paste the text instead.",
            )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Could not read '{filename}': {e}")

    if not text.strip():
        raise HTTPException(400, f"No text found in '{filename}' — file may be empty or image-only.")

    return {"status": "ok", "filename": filename, "text": text}


# ── Endpoints ──────────────────────────────────────────────────────────────────

@router.post("/raw")
async def ingest_raw(req: RawIngestRequest):
    """
    Parse raw pasted text with Claude and return a structured preview.
    Does NOT write to the database.
    """
    if req.type not in ("fighter_profile", "fight_history"):
        raise HTTPException(400, "type must be 'fighter_profile' or 'fight_history'")
    if not req.raw_text.strip():
        raise HTTPException(400, "raw_text cannot be empty")
    if req.type == "fight_history" and not req.fighter_name:
        raise HTTPException(400, "fighter_name is required for fight_history ingestion")

    prompt = (
        RAW_FIGHTER_PROFILE_PROMPT
        if req.type == "fighter_profile"
        else RAW_FIGHT_HISTORY_PROMPT
    )

    user_content = (
        f"Fighter: {req.fighter_name}\n\n" if req.fighter_name else ""
    ) + f"Raw source text:\n{req.raw_text[:12000]}"

    claude = get_claude()
    try:
        parsed = claude.analyze(prompt, user_content)
    except ClaudeOverloadedError as e:
        return {
            "status": "error",
            "error": "The AI service is temporarily overloaded. Please wait a few seconds and try again.",
            "detail": str(e),
            "retryable": True,
        }
    except ClaudeEmptyResponseError as e:
        return {
            "status": "error",
            "error": "Could not parse fighter data from the pasted text. Make sure you copied the full page and try again.",
            "detail": str(e),
            "retryable": False,
        }
    except Exception as e:
        return {
            "status": "error",
            "error": f"Unexpected error: {str(e)}",
        }

    # ── Normalise fight_history: Claude may return a single dict for Format B ──
    if req.type == "fight_history" and isinstance(parsed, dict):
        parsed = [parsed]

    # ── Post-process fight_history: resolve event dates from events table ──────
    if req.type == "fight_history" and isinstance(parsed, list):
        parsed = _enrich_bout_dates(parsed)
        # Sort by event_date ASC; bouts with no date go to end
        parsed = sorted(
            parsed,
            key=lambda b: (b.get("event_date") is None, b.get("event_date") or "")
        )

    # ── OpenAI audit — compare raw text vs Claude output ──────────────────────
    # Runs asynchronously-ish (in the same request) but never blocks approval.
    # Only audit fight_history — that's where Win/Loss flips happen.
    audit = {"status": "skipped", "issues": []}
    if req.type == "fight_history" and parsed:
        try:
            audit = audit_parse(req.raw_text, parsed)
        except Exception as audit_err:
            print(f"[Audit] Non-fatal error: {audit_err}")
            audit = {"status": "error", "issues": [], "error": str(audit_err)}

    # ── Validation preview (don't block — let operator see what's missing) ────
    if req.type == "fighter_profile":
        errors = validate_fighter_profile(parsed if isinstance(parsed, dict) else {})
    else:
        errors = validate_fight_history(parsed if isinstance(parsed, list) else [])

    return {
        "status": "preview",
        "type": req.type,
        "data": parsed,
        "raw_text": req.raw_text,          # echoed back so frontend can show snippets
        "audit": audit,                    # OpenAI audit result
        "validation_errors": errors,
        "ready": len(errors) == 0,
        "fighter_name": req.fighter_name,
    }


@router.post("/approve")
async def ingest_approve(req: ApproveRequest):
    """
    Validate + save approved data to Supabase for review.
    Final publishing to GRIT happens from the review dashboard approval step.
    """
    supabase = get_supabase()

    if req.type == "fighter_profile":
        if not isinstance(req.data, dict):
            raise HTTPException(400, "data must be an object for fighter_profile")

        errors = validate_fighter_profile(req.data)
        if errors:
            return {"status": "error", "validation_errors": errors}

        upsert = _build_fighter_upsert(req.data)
        first_name = upsert["first_name"]
        last_name  = upsert["last_name"]

        existing = find_fighter_by_name(supabase, first_name, last_name)
        if existing:
            fighter_id = existing["id"]
            update_payload = {k: v for k, v in upsert.items()}
            supabase.table("fighters").update(update_payload).eq("id", fighter_id).execute()
            action = "updated"
        else:
            res = supabase.table("fighters").insert(upsert).execute()
            fighter_id = res.data[0]["id"]
            action = "created"

        print(f"[Ingest] Fighter {action}: {first_name} {last_name} ({fighter_id})")

        # Queued for human review; dashboard approval performs the GRIT push.
        push_status = "review_required"

        return {
            "status": "success",
            "action": action,
            "fighter_id": fighter_id,
            "name": f"{first_name} {last_name}",
            "review_status": "pending",
            "push_status": push_status,
        }

    elif req.type == "fight_history":
        if not req.fighter_id:
            raise HTTPException(400, "fighter_id is required for fight_history approval")
        if not isinstance(req.data, list):
            raise HTTPException(400, "data must be an array of bouts for fight_history")

        errors = validate_fight_history(req.data)
        if errors:
            return {"status": "error", "validation_errors": errors}

        # Verify the fighter exists
        fighter_res = supabase.table("fighters").select("id,first_name,last_name")\
            .eq("id", req.fighter_id).execute()
        if not fighter_res.data:
            raise HTTPException(404, f"Fighter ID {req.fighter_id} not found")
        fighter = fighter_res.data[0]

        # ── BLOCKER 1: Duplicate prevention ──────────────────────────────────
        # Build a set of (opponent_name_lower, event_name_lower) combos
        # already in the DB for this fighter — skip any incoming bouts that match.
        existing_keys: set[tuple[str, str]] = set()
        try:
            ex_res = supabase.table("fight_history")\
                .select("opponent_name,event_name")\
                .eq("fighter_id", req.fighter_id)\
                .execute()
            for row in (ex_res.data or []):
                opp  = (row.get("opponent_name") or "").strip().lower()
                evt  = _normalize_event_name(row.get("event_name") or "").lower()
                existing_keys.add((opp, evt))
        except Exception as e:
            print(f"[Ingest] Could not fetch existing bouts for dedup (continuing): {e}")

        # Build a local event-name → event_id cache to avoid N+1 queries
        _event_id_cache: dict[str, str] = {}
        NIL_UUID = "00000000-0000-0000-0000-000000000000"

        def _resolve_event_id(raw_name: str) -> str:
            """Look up event UUID by normalised name; fall back to nil UUID."""
            norm = _normalize_event_name(raw_name)
            if not norm:
                return NIL_UUID
            if norm in _event_id_cache:
                return _event_id_cache[norm]
            try:
                res = supabase.table("events").select("id,name").ilike("name", f"{norm}%").limit(1).execute()
                eid = res.data[0]["id"] if res.data else NIL_UUID
            except Exception:
                eid = NIL_UUID
            _event_id_cache[norm] = eid
            return eid

        rows = []
        skipped = 0
        fighters_created = 0
        for idx, bout in enumerate(req.data):
            if not isinstance(bout, dict):
                continue
            event_name = _normalize_event_name(bout.get("event_name") or "")
            opponent   = (bout.get("opponent_name") or "").strip()
            dedup_key  = (opponent.lower(), event_name.lower())

            if dedup_key in existing_keys:
                skipped += 1
                print(f"[Ingest] Skipping duplicate bout: {opponent} @ {event_name}")
                continue

            # Resolve opponent_id — look up existing fighter or create a shell.
            # Non-blocking: if lookup/creation fails, opponent_id stays None.
            opponent_id, was_created = _resolve_or_create_opponent(supabase, opponent)
            if was_created:
                fighters_created += 1

            title_fight = bool(bout.get("title_fight", False))
            event_date  = bout.get("event_date")
            event_id    = _resolve_event_id(bout.get("event_name") or "")
            rows.append({
                "fighter_id":            req.fighter_id,
                "event_id":              event_id,
                "opponent_name":         opponent,
                "opponent_id":           opponent_id,
                "event_name":            event_name,
                "event_date":            event_date,
                # ── NOT NULL DB columns — always provide a value to avoid 23502 ──
                # Text columns: empty string satisfies NOT NULL; real value preferred.
                "result":                bout.get("result") or "",
                "method":                bout.get("method") or "",
                "time":                  bout.get("time") or "",
                # Int columns: 0 = unknown/not applicable.
                "round":                 bout.get("round") or 0,
                "fight_duration_seconds": bout.get("fight_duration_seconds") or 0,
                # JSONB column: empty object satisfies NOT NULL.
                "location":              bout.get("location") or {},
                # fight_type / bout_order — use agent3_history defaults.
                "fight_type":            bout.get("fight_type") or ("Title Bout" if title_fight else "Bout"),
                "bout_order":            bout.get("bout_order") or (idx + 1),
                # ── Nullable columns (null allowed) ────────────────────────────
                "title_fight":           title_fight,
                "method_detail":         bout.get("method_detail"),
                "rounds_scheduled":      bout.get("rounds_scheduled"),
                "referee":               bout.get("referee"),
                "stats":                 bout.get("bout_stats"),
            })

        if not rows and skipped == 0:
            return {"status": "error", "validation_errors": ["No valid bouts to insert"]}

        fights_created_count = 0
        if rows:
            # upsert with ignore_duplicates=True → INSERT … ON CONFLICT DO NOTHING
            # Requires the unique index on (fighter_id, event_name, opponent_name)
            # created by migrations 004 / 009 / 013.
            # opponent_id requires migration 014 — fall back without it if the
            # column doesn't exist yet (safe to run before the migration is applied).
            def _strip_opponent_id(r_list):
                return [{k: v for k, v in r.items() if k != "opponent_id"} for r in r_list]

            try:
                supabase.table("fight_history").upsert(
                    rows,
                    on_conflict="fighter_id,event_name,opponent_name",
                    ignore_duplicates=True,
                ).execute()
            except Exception as upsert_err:
                err = str(upsert_err)
                rows_no_opp = _strip_opponent_id(rows)

                if "42P10" in err or "no unique or exclusion constraint" in err:
                    # Dedup index not yet created — application-level pre-fetch
                    # dedup already filtered duplicates, so plain insert is safe.
                    # Run migrations 004/009/013 to enable DB-level dedup.
                    print(
                        "[Ingest] Dedup unique index missing on fight_history — "
                        "run migrations 004/009/013 to enable ON CONFLICT. "
                        "Falling back to plain INSERT (application-level dedup still active)."
                    )
                    try:
                        supabase.table("fight_history").insert(rows).execute()
                    except Exception as insert_err2:
                        if "opponent_id" in str(insert_err2):
                            print("[Ingest] opponent_id column also missing — retrying without it.")
                            supabase.table("fight_history").insert(rows_no_opp).execute()
                        else:
                            raise

                elif "opponent_id" in err:
                    # opponent_id column missing (migration 014 not yet applied) —
                    # retry upsert without opponent_id.
                    print(
                        "[Ingest] opponent_id column missing — "
                        "run migration 014_fight_history_opponent_id.sql. "
                        "Retrying upsert without opponent_id."
                    )
                    try:
                        supabase.table("fight_history").upsert(
                            rows_no_opp,
                            on_conflict="fighter_id,event_name,opponent_name",
                            ignore_duplicates=True,
                        ).execute()
                    except Exception as upsert_err2:
                        if "42P10" in str(upsert_err2) or "no unique or exclusion constraint" in str(upsert_err2):
                            print("[Ingest] Dedup index also missing — falling back to plain INSERT.")
                            supabase.table("fight_history").insert(rows_no_opp).execute()
                        else:
                            raise
                else:
                    raise
            fights_created_count = len(rows)
            print(
                f"[Ingest] fights_created={fights_created_count} "
                f"fights_skipped={skipped} fighters_created={fighters_created} "
                f"for {fighter['first_name']} {fighter['last_name']} ({req.fighter_id})"
            )
        else:
            print(
                f"[Ingest] All {skipped} bouts already exist for "
                f"{fighter['first_name']} {fighter['last_name']} — nothing inserted"
            )

        # Queued for human review; dashboard approval performs the GRIT push.
        push_status = "review_required"

        return {
            "status":           "success",
            "fights_created":   fights_created_count,
            "fights_skipped":   skipped,
            "fighters_created": fighters_created,
            "fighter_id":       req.fighter_id,
            "review_status":    "pending",
            "push_status":      push_status,
        }

    else:
        raise HTTPException(400, "type must be 'fighter_profile' or 'fight_history'")


@router.post("/check-names")
async def check_names(names: list[str]):
    """Check a list of fighter names against the DB to see who already exists."""
    supabase = get_supabase()
    results = []
    for full_name in names:
        if not full_name.strip():
            continue
        # normalize_name: unicode → ASCII, lowercase, collapse spaces, strip punctuation
        first, last = normalize_name(full_name)
        existing = find_fighter_by_name(supabase, first, last)
        results.append({
            "name": full_name,
            "exists": existing is not None,
            "id": existing["id"] if existing else None
        })
    return results

@router.get("/fighters/search")
async def search_fighters(q: str = ""):
    """Quick fighter search for the fight_history flow (select who the history belongs to)."""
    supabase = get_supabase()
    if not q.strip():
        res = supabase.table("fighters")\
            .select("id,first_name,last_name,weight_class,wins,losses")\
            .order("last_name")\
            .limit(50)\
            .execute()
    else:
        parts = q.strip().lower().split()
        first = parts[0] if parts else ""
        last  = parts[1] if len(parts) > 1 else ""
        if last:
            res = supabase.table("fighters")\
                .select("id,first_name,last_name,weight_class,wins,losses")\
                .ilike("last_name", f"%{last}%")\
                .limit(20)\
                .execute()
        else:
            import re
            # Search both first and last name
            res_first = supabase.table("fighters")\
                .select("id,first_name,last_name,weight_class,wins,losses")\
                .ilike("first_name", f"%{first}%")\
                .limit(10)\
                .execute()
            res_last = supabase.table("fighters")\
                .select("id,first_name,last_name,weight_class,wins,losses")\
                .ilike("last_name", f"%{first}%")\
                .limit(10)\
                .execute()
            combined = {r["id"]: r for r in (res_first.data + res_last.data)}
            return list(combined.values())

    return res.data
